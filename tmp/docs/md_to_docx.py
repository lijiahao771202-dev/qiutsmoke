from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys

md_path = "/Users/lijiahao/.gemini/antigravity/brain/5d01db91-1ed2-4ca8-bf0d-c56984b9cf2f/surge_architecture.md"
docx_path = "/Users/lijiahao/.gemini/antigravity/brain/5d01db91-1ed2-4ca8-bf0d-c56984b9cf2f/surge_architecture.docx"

doc = Document()

# Configure basic style
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

try:
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
except Exception as e:
    print(f"Error reading file: {e}")
    sys.exit(1)

# Split by blocks properly
blocks = content.split('\n\n')

in_code_block = False
code_text = ""

for block in blocks:
    if block.startswith('```') and not in_code_block:
        in_code_block = True
        # remove the opening tag
        block = block.split('\n', 1)[1] if '\n' in block else ""
        
    if in_code_block:
        if '```' in block:
            in_code_block = False
            code_text += '\n\n' + block.replace('```', '')
            p = doc.add_paragraph(code_text.strip())
            p.style.font.name = 'Courier'
            p.style.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
            code_text = ""
        else:
            if code_text != "":
                 code_text += '\n\n' + block
            else:
                 code_text = block
        continue
        
    # Headers
    if block.startswith('# '):
        heading = doc.add_heading(block.replace('# ', '').strip(), level=1)
    elif block.startswith('## '):
        heading = doc.add_heading(block.replace('## ', '').strip(), level=2)
    elif block.startswith('### '):
        heading = doc.add_heading(block.replace('### ', '').strip(), level=3)
        
    # Lists
    elif block.startswith('- ') or block.startswith('1. ') or block.startswith('2. ') or block.startswith('3. '):
        lines = block.split('\n')
        for line in lines:
            line = line.strip().replace('**', '')
            if line.startswith('- '):
                doc.add_paragraph(line.replace('- ', '', 1), style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. '):
                doc.add_paragraph(line.replace('1. ', '', 1).replace('2. ', '', 1).replace('3. ', '', 1), style='List Number')
            elif line != "":
                doc.add_paragraph(line)
                
    # Separator
    elif block.startswith('---'):
        p = doc.add_paragraph('----------------------------------------------------')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    # Normal text
    else:
        clean_text = block.replace('**', '')
        if clean_text.strip() != "":
            doc.add_paragraph(clean_text)

doc.save(docx_path)
print(f"Successfully generated {docx_path}")
